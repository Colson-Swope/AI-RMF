from docx import Document
from datetime import datetime


def create_docx(input_file, output_file):
    doc = Document()
    
    title = "Operating System Patch Management RMF Compliance"

    doc.add_heading(title, 2)

    with open(input_file) as f:
        for line in f:
            doc.add_paragraph(line)

    doc.save(output_file)
