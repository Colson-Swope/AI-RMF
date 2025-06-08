from docx import Document
from datetime import datetime


def create_docx(input_file, output_file):
    doc = Document()

    # Add title and timestamp
    doc.add_heading("Operating System Patch Management RMF Compliance", level=1)
    timestamp = datetime.now().strftime("Created %B %d, %Y at %H:%M:%S")
    doc.add_paragraph(timestamp, style='Normal')
    doc.add_paragraph("")  # spacer

    # Parse input file
    with open(input_file, 'r') as f:
        lines = f.readlines()

    bullet_buffer = []

    def flush_bullets():
        for item in bullet_buffer:
            doc.add_paragraph(item, style='List Bullet')
        bullet_buffer.clear()

    # Build output file
    for line in lines:
        line = line.strip()
        if line.startswith(('- ', '* ', '+ ')):
            bullet_buffer.append(line[2:])
        else:
            if bullet_buffer:
                flush_bullets()
            if line.startswith('***') and line.endswith('***'):
                doc.add_heading(line[3:-3].strip(), level=2)
            elif line.startswith('**') and line.endswith('**'):
                doc.add_heading(line[2:-2].strip(), level=2)
            elif line:
                doc.add_paragraph(line, style='Normal')

    if bullet_buffer:
        flush_bullets()

    # Save the document
    doc.save(output_file)

