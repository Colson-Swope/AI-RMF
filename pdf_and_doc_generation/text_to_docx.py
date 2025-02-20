from docx import Document
from datetime import datetime


def create_docx(input_file, output_file):
    doc = Document()
    # pdf.set_auto_page_break(True, margin=margin_size)

    title = "AI RMF Report"
    date = str(datetime.now())

    doc.add_heading(title, 0)
    doc.add_heading(date, 1)


    with open(input_file) as f:
        for line in f:
            doc.add_paragraph(line)

    doc.save(output_file)


