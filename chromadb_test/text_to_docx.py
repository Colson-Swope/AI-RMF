from docx import Document
from docx.shared import Pt
from datetime import datetime


def create_docx(input_file, output_file):
    doc = Document()

    title = "Operating System Patch Management RMF Compliance"
    doc.add_heading(title, level=1)

    doc.add_paragraph(f"Created {datetime.now().strftime('%B %d, %Y at %H:%M:%S')}")
    doc.add_paragraph("")

    with open(input_file, 'r') as f:
        lines = f.readlines()

    current_section = None
    buffer = []

    def add_section(title, content):
        if title:
            doc.add_heading(title, level=2)
        for line in content:
            line = line.strip()
            if line.startswith('- '):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)

    for line in lines:
        line = line.strip()
        if line.startswith('***') and line.endswith('***'):
            if current_section and buffer:
                add_section(current_section, buffer)
                buffer = []
            current_section = line.strip('* ').strip()
        else:
            buffer.append(line)

    if current_section and buffer:
        add_section(current_section, buffer)

    doc.save(output_file)
